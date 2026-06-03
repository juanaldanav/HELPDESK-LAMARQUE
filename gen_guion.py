# -*- coding: utf-8 -*-
# Regenera GUION_DEMO.docx actualizado. Uso: python soporte-v2/gen_guion.py
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = r"D:\LAMARQUE\GLPI\soporte-v2\entregables"
os.makedirs(OUT, exist_ok=True)
BRAND = RGBColor(0x00, 0x69, 0x70); INK = RGBColor(0x0f, 0x1c, 0x1d); MUT = RGBColor(0x5d, 0x6f, 0x70)
URL = "https://soporte.lamarque.mx"

g = Document()
st = g.styles['Normal']; st.font.name = 'Calibri'; st.font.size = Pt(11); st.font.color.rgb = INK

def h1(t):
    p = g.add_paragraph(); r = p.add_run(t); r.bold = True; r.font.size = Pt(20); r.font.color.rgb = BRAND
def h2(t):
    p = g.add_paragraph(); p.space_before = Pt(8); r = p.add_run(t); r.bold = True; r.font.size = Pt(15); r.font.color.rgb = BRAND
def h3(t):
    p = g.add_paragraph(); r = p.add_run(t); r.bold = True; r.font.size = Pt(12.5); r.font.color.rgb = INK
def p(t, size=11, color=INK, bold=False, italic=False):
    par = g.add_paragraph(); r = par.add_run(t); r.font.size = Pt(size); r.font.color.rgb = color; r.bold = bold; r.italic = italic; return par
def b(t):  # viñeta
    par = g.add_paragraph(style='List Bullet'); r = par.add_run(t); r.font.size = Pt(11)
def step(n, t):
    par = g.add_paragraph(); r = par.add_run(f"{n}. "); r.bold = True; r.font.color.rgb = BRAND; r.font.size = Pt(11)
    r2 = par.add_run(t); r2.font.size = Pt(11)

h1("Soporte Lamarque — Guion de demostración")
p("Portal de mantenimiento y activos para las 16 sucursales. Interfaz a la medida sobre la base de datos de GLPI (motor robusto) con experiencia simple por rol.", 11)
p(URL + "   ·   la página principal abre directo el portal.", 10, BRAND, bold=True)

h2("Preparación (1 min)")
b("Ten a la mano: un celular (rol Sucursal y rol Técnico) y una laptop (rol Dalia / Mejora Continua).")
b("Cuentas demo — contraseña Lamarque2026!:  guadalupe (sucursal) · alfonso / noe / armando (técnico) · dalia (Mejora Continua).")
b("Los correos están en modo pruebas: todas las notificaciones llegan a sistemas@lamarque.mx (no se molesta a las sucursales durante la demo).")
b("Mensaje de apertura: «Es el mismo motor de datos de GLPI, pero una interfaz hecha a la medida de una cafetería: cada quien ve solo lo que necesita.»")

h2("Acto 1 — La sucursal reporta (celular, ~2 min)")
step(1, "Entra como guadalupe. Muestra el Inicio: saludo, botón grande “Reportar un problema”, contadores (abiertos / en espera / resueltos) y los tickets recientes. Abajo, la barra de pestañas Inicio · Tickets · Activos.")
step(2, "Toca “Reportar un problema”. Llena: título, categoría, urgencia (botones de color), descripción.")
step(3, "En “Equipo afectado” busca por folio o nombre (ej. LMQ-GDL-REF-0005). Recalca: solo aparecen activos de SU sucursal.")
step(4, "Adjunta una o varias fotos (hasta 5) y envía. Se abre el hilo del ticket tipo chat con el acuse de recibo.")
step(5, "Entra a “Mis activos”: cuadros por categoría con su conteo. Abre uno y muestra “Ver referencia” → abre la foto del equipo en Drive (sirve para ubicarlo físicamente).")

h2("Acto 2 — Dalia coordina (laptop, ~3 min)")
step(1, "Entra como dalia. Dashboard: tickets del periodo, completados, “Pendientes (hoy)” = abiertos de cualquier mes, correctivo vs preventivo, top sucursales, categorías y atenciones por proveedor (interno vs externos).")
step(2, "Cambia el “Periodo” a un mes con historial (ej. mayo) para ver las gráficas llenas. Explica que es el análisis que antes hacía en Excel, ahora en vivo.")
step(3, "Ve a Tickets: consolidado de TODAS las sucursales. Muestra filtros (sucursal, categoría, técnico, estado, fechas), el selector “Mostrar 25/50/100” con paginación, y “Exportar CSV”.")
step(4, "Abre el ticket recién creado. En “Asignar / actualizar” muestra las dos formas de atender:")
b("Técnico interno: elige a un técnico, urgencia y una Fecha de atención → se notifica y aparece en el Calendario.")
b("Proveedor externo: cambia a “Proveedor externo”, elige uno del catálogo o crea uno nuevo (“+ Nuevo proveedor”).")
step(5, "Escribe un comentario en el ticket (Dalia también chatea con sucursal y técnico).")
step(6, "Abre el Calendario: el mantenimiento quedó agendado, con colores por tipo. Haz clic en cualquier día → se abre “Agendar mantenimiento” con esa fecha lista (sucursal + tipo + técnico).")

h2("Acto 3 — El técnico atiende (celular / PWA, ~2 min)")
step(1, "Entra como el técnico asignado. Si es la demo de instalación: en el navegador del celular, menú → “Agregar a pantalla de inicio” → queda como app.")
step(2, "“Mis tareas”: muestra los tickets asignados (filtro Pendientes / En curso) Y la sección “Mantenimientos programados” — los preventivos/correctivos del calendario le llegan aquí, no solo por correo.")
step(3, "Abre la tarea: ve la sucursal, el equipo vinculado (con “Ver referencia”), el hilo; puede comentar y adjuntar foto.")
step(4, "Toca “Cerrar con hoja de servicio”: marca el checklist, escribe observaciones, sube foto de evidencia y cierra.")
step(5, "Al cerrar se genera la HOJA DE SERVICIO EN PDF (membretada con el logo) y se envía por correo a la sucursal y a Dalia.")

h2("Acto 4 — Cierre del ciclo y administración (~2 min)")
step(1, "Regresa a la sucursal: el ticket aparece Cerrado, con la hoja en su hilo y botón para descargar el PDF.")
step(2, "Aislamiento: entra con otra sucursal y muestra que NO ve los tickets ni activos de Guadalupe.")
step(3, "Proveedor externo: en un ticket atendido por proveedor, Dalia usa “Cerrar ticket (coordinación)” (el proveedor no entra al sistema) — también genera la hoja PDF.")
step(4, "Equipo (panel de Dalia): buscador + tarjetas por usuario. Demuestra: crear usuario, cambiar rol y sucursal, editar nombre/correo, activar/desactivar, restablecer contraseña y eliminar. (ti.admin queda protegido.)")
step(5, "Mi cuenta: cualquier usuario cambia su propia contraseña (icono de persona arriba). Útil al pasar de datos de prueba a reales.")

h2("Puntos clave para el cliente")
b("Mismo motor robusto de GLPI (3,000+ activos, 1,000+ tickets históricos) con interfaz simple y propia.")
b("Cada rol ve solo lo necesario: sucursal reporta, técnico atiende, Dalia coordina y mide.")
b("Móvil para sucursal y técnico (instalable como app); escritorio para Dalia. Todo responsivo.")
b("Trazabilidad completa: foto del reporte, hilo de conversación, hoja de servicio en PDF y notificación por correo.")
b("Mantenimiento preventivo (calendario, 83 eventos del cronograma) y correctivo (tickets) en un solo lugar.")
b("Atención interna (técnicos) o externa (proveedores), medida en el dashboard sin exponer costos.")

h2("Notas para quien presenta")
b("Correos en modo pruebas → llegan a sistemas@lamarque.mx. Para producción se activa el envío real a sucursales.")
b("Los técnicos no necesitan correo: sus tareas les llegan en la app y su gestión (clave, rol) la hace Dalia o TI.")
b("Si una gráfica del dashboard sale vacía, es porque ese mes no tiene tickets: cambia el periodo.")

g.save(os.path.join(OUT, "GUION_DEMO.docx"))
print("GUION_DEMO.docx actualizado en", OUT)
