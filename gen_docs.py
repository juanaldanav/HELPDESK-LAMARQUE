# -*- coding: utf-8 -*-
# Genera GUION_DEMO.docx y MANUAL_USUARIO.docx con capturas. Uso: python soporte-v2/gen_docs.py
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE = r"D:\LAMARQUE\GLPI"
OUT = r"D:\LAMARQUE\GLPI\soporte-v2\entregables"
os.makedirs(OUT, exist_ok=True)
BRAND = RGBColor(0x00, 0x69, 0x70)
INK = RGBColor(0x0f, 0x1c, 0x1d)
MUT = RGBColor(0x5d, 0x6f, 0x70)
URL = "https://soporte.lamarque.mx/soporte-v2/"

def img(path):
    return os.path.join(BASE, path)

def base_style(doc):
    st = doc.styles['Normal']
    st.font.name = 'Calibri'; st.font.size = Pt(11); st.font.color.rgb = INK

def h1(doc, text):
    p = doc.add_paragraph(); r = p.add_run(text); r.bold = True; r.font.size = Pt(20); r.font.color.rgb = BRAND
    p.space_after = Pt(6); return p

def h2(doc, text):
    p = doc.add_paragraph(); r = p.add_run(text); r.bold = True; r.font.size = Pt(15); r.font.color.rgb = BRAND
    return p

def para(doc, text, size=11, color=INK, bold=False):
    p = doc.add_paragraph(); r = p.add_run(text); r.font.size = Pt(size); r.font.color.rgb = color; r.bold = bold
    return p

def shot(doc, file, w=5.2):
    fp = img(file)
    if os.path.isfile(fp):
        doc.add_picture(fp, width=Inches(w))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

def caption(doc, text):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.italic = True; r.font.size = Pt(9); r.font.color.rgb = MUT

# ───────────────────────── GUION DE DEMOSTRACIÓN ─────────────────────────
g = Document(); base_style(g)
h1(g, "Soporte Lamarque — Guion de demostración")
para(g, "Portal de mantenimiento y activos · 16 sucursales. App a la medida sobre la base de datos de GLPI.", 11, MUT)
para(g, URL, 10, BRAND)
para(g, "Cuentas demo (contraseña Lamarque2026!): guadalupe (sucursal) · alfonso/noe/armando (técnico) · dalia (Mejora Continua).", 10, MUT)

h2(g, "Antes de empezar")
for t in ["Ten 3 pestañas o 1 móvil + 1 laptop: una por rol.",
          "El correo de notificaciones está en modo pruebas: todo llega a sistemas@lamarque.mx (no se molesta a sucursales).",
          "Duración sugerida: 8–10 minutos."]:
    para(g, "•  " + t)

steps = [
    ("1. La sucursal reporta (celular)", [
        "Entra como guadalupe. Muestra el inicio: saludo, botón grande “Reportar un problema”, contadores y tickets recientes.",
        "Toca “Reportar un problema”. Llena: título “La vitrina no enfría”, categoría Refrigeración, urgencia (botones de color), descripción.",
        "Busca el equipo por folio (ej. LMQ-GDL-REF-0005) — el buscador filtra solo activos de TU sucursal.",
        "Adjunta una o varias fotos (hasta 5) y envía. Aparece el hilo del ticket con el acuse.",
    ]),
    ("2. Dalia coordina (escritorio)", [
        "Entra como dalia. Dashboard: KPIs del mes, top sucursales, correctivo vs preventivo, categorías y proveedores.",
        "Ve a Tickets: consolidado de TODAS las sucursales con filtros (sucursal, categoría, técnico, estado, fechas) y exportación CSV.",
        "Abre el ticket recién creado. Asigna un técnico, fija urgencia y una FECHA DE ATENCIÓN → se notifica y aparece en el Calendario.",
        "Abre Calendario: el mantenimiento queda agendado (preventivos del cronograma + correctivos con fecha), con colores por tipo.",
    ]),
    ("3. El técnico atiende (celular / PWA)", [
        "Entra como el técnico asignado. “Mis tareas” muestra solo sus tickets, ordenados por urgencia.",
        "Abre el ticket: ve el activo vinculado con botón “Ver referencia” (foto del equipo en Drive), el hilo y puede comentar/adjuntar.",
        "Toca “Cerrar con hoja de servicio”: marca el checklist, escribe observaciones, sube foto de evidencia y cierra.",
        "Al cerrar se genera la HOJA DE SERVICIO EN PDF y se envía por correo (con la marca) a la sucursal y a Dalia.",
    ]),
    ("4. Cierre del ciclo", [
        "Vuelve a la sucursal: el ticket aparece como Cerrado con la hoja de servicio en su hilo y botón para descargar el PDF.",
        "Muestra el aislamiento: otra sucursal NO ve ese ticket ni esos activos.",
        "En Equipo (Dalia): crea/edita usuarios, cambia roles y sucursal, activa/desactiva y restablece contraseñas.",
    ]),
]
for title, items in steps:
    h2(g, title)
    for it in items:
        para(g, "•  " + it)

h2(g, "Mensajes clave para el cliente")
for t in ["Mismo motor de datos robusto de GLPI, pero una interfaz simple hecha a la medida de una cafetería.",
          "Cada rol ve solo lo que necesita: sin ruido, sin módulos corporativos.",
          "Móvil para sucursal y técnico; escritorio para Dalia. Todo responsivo.",
          "Trazabilidad completa: foto del reporte, hilo, hoja de servicio en PDF y correo."]:
    para(g, "•  " + t)

g.save(os.path.join(OUT, "GUION_DEMO.docx"))
print("GUION_DEMO.docx OK")

# ───────────────────────── MANUAL DE USUARIO ─────────────────────────
m = Document(); base_style(m)
h1(m, "Soporte Lamarque — Manual de usuario")
para(m, "Portal de mantenimiento y activos. Guía para Sucursal, Técnico y Mejora Continua (Dalia).", 11, MUT)
para(m, URL, 10, BRAND)

h2(m, "Índice")
for t in ["1. Acceso al sistema (todos)",
          "2. Sucursal — reportar y dar seguimiento",
          "3. Técnico — atender y cerrar con hoja de servicio",
          "4. Mejora Continua (Dalia) — coordinar, medir y administrar",
          "5. Preguntas frecuentes"]:
    para(m, t, 11, INK, bold=True)

m.add_page_break()

# 1. Acceso
h2(m, "1. Acceso al sistema")
para(m, "Entra desde el navegador del celular o computadora a:", 11)
para(m, URL, 11, BRAND, bold=True)
para(m, "Escribe tu usuario y contraseña. Si la olvidaste, usa “¿Olvidaste tu contraseña?”: recibirás un enlace por correo (válido 15 minutos) para crear una nueva.", 11)
shot(m, "man-login.png", 2.6); caption(m, "Pantalla de acceso")

# 2. Sucursal
m.add_page_break(); h2(m, "2. Sucursal — reportar y dar seguimiento")
para(m, "Inicio. Verás un saludo, el botón “Reportar un problema”, tus contadores (abiertos, en espera, resueltos) y tus tickets recientes. Abajo, la barra con Inicio · Tickets · Activos.", 11)
shot(m, "man-suc-home.png", 2.6); caption(m, "Inicio de sucursal")
para(m, "Reportar un problema. Escribe qué pasó, elige la categoría, la urgencia (botones de color), describe el problema, opcionalmente busca el equipo por folio o nombre, y adjunta hasta 5 fotos. Toca “Enviar reporte”.", 11)
shot(m, "s-new.png", 2.6); caption(m, "Formulario de reporte")
para(m, "Seguimiento (hilo). Cada ticket es una conversación: ves el reporte inicial, las respuestas del técnico y el estado. Puedes escribir mensajes y adjuntar fotos mientras esté abierto. Si tiene un equipo vinculado, “Ver referencia” abre su foto.", 11)
shot(m, "man-suc-thread.png", 2.6); caption(m, "Hilo del ticket")
para(m, "Mis activos. Tus equipos por categoría (con su conteo). Toca una categoría para ver el detalle; cada equipo muestra su folio, estado y “Ver referencia”.", 11)
shot(m, "man-suc-assets.png", 2.6); caption(m, "Activos de la sucursal por categoría")

# 3. Técnico
m.add_page_break(); h2(m, "3. Técnico — atender y cerrar")
para(m, "Mis tareas. Lista de tickets asignados a ti, ordenados por urgencia. Usa el filtro Pendientes / En curso.", 11)
shot(m, "man-tec-home.png", 2.6); caption(m, "Tareas del técnico")
para(m, "Detalle de la tarea. Ves la sucursal, el equipo vinculado (con “Ver referencia”), el hilo y puedes comentar o adjuntar fotos.", 11)
shot(m, "man-tec-detail.png", 2.6); caption(m, "Detalle de la tarea")
para(m, "Cerrar con hoja de servicio. Toca el botón y completa el checklist (limpieza, cableado, tornillería, lubricación, etc.), escribe observaciones y sube una foto de evidencia. Al cerrar, se genera la hoja de servicio en PDF y se envía por correo a la sucursal y a Dalia.", 11)
shot(m, "man-tec-sheet.png", 2.6); caption(m, "Hoja de servicio (checklist)")
shot(m, "man-pdf.png", 2.6); caption(m, "Hoja de servicio generada en PDF")

# 4. Dalia
m.add_page_break(); h2(m, "4. Mejora Continua (Dalia)")
para(m, "Dashboard. Indicadores del periodo: total de tickets, completados, pendientes, correctivo vs preventivo, top sucursales, categorías más recurrentes y atenciones por proveedor (interno vs externo). Cambia el mes arriba a la derecha.", 11)
shot(m, "man-dalia-dashboard.png", 6.2); caption(m, "Dashboard de indicadores")
para(m, "Tickets (consolidado). Todas las sucursales en una tabla con filtros por sucursal, categoría, técnico, estado y rango de fechas. Exporta a CSV. Haz clic en una fila para abrir el ticket.", 11)
shot(m, "man-dalia-tickets.png", 6.2); caption(m, "Consolidado de tickets")
para(m, "Asignar técnico. Dentro de un ticket: elige técnico, urgencia y fecha de atención (se notifica y se agenda). Botón Imprimir/PDF disponible.", 11)
shot(m, "dd-view.png", 6.2); caption(m, "Ticket: asignación y seguimiento")
para(m, "Calendario. Vista mensual de mantenimientos preventivos (del cronograma) y correctivos (con fecha). Colores por tipo. Botón “Agendar mantenimiento” para crear uno nuevo (sucursal, tipo, fecha/rango, técnico) que se notifica por correo.", 11)
shot(m, "man-dalia-calendar.png", 6.2); caption(m, "Calendario de mantenimientos")
para(m, "Activos (consolidado). Inventario de todas las sucursales filtrable por sucursal y categoría. Permite dar de baja un equipo (queda en histórico).", 11)
shot(m, "man-dalia-assets.png", 6.2); caption(m, "Inventario consolidado")
para(m, "Equipo (usuarios y roles). Crea usuarios, cambia su rol (Sucursal / Técnico / Mejora Continua) y su sucursal, actívalos o desactívalos y restablece contraseñas. Las cuentas de administración del sistema están protegidas.", 11)
shot(m, "man-dalia-users.png", 6.2); caption(m, "Gestión de equipo y roles")

# 5. FAQ
m.add_page_break(); h2(m, "5. Preguntas frecuentes")
faqs = [
    ("¿Puedo ver tickets de otra sucursal?", "No. Cada sucursal ve únicamente sus propios tickets y activos. Solo Mejora Continua ve todo."),
    ("Olvidé mi contraseña.", "Usa “¿Olvidaste tu contraseña?” en el acceso. Llega un enlace al correo, válido 15 minutos."),
    ("¿Para qué sirve “Ver referencia”?", "Abre la foto del equipo (Google Drive) para ubicarlo físicamente y confirmar de cuál se trata."),
    ("¿Quién recibe la hoja de servicio?", "Al cerrar el ticket, el PDF se envía por correo a la sucursal y a Mejora Continua, y queda disponible en el portal."),
    ("¿Funciona en celular?", "Sí. Sucursal y técnico están pensados para celular; Dalia para computadora. Todo se adapta a la pantalla."),
]
for q, a in faqs:
    para(m, q, 11, BRAND, bold=True); para(m, a, 11)

m.save(os.path.join(OUT, "MANUAL_USUARIO.docx"))
print("MANUAL_USUARIO.docx OK")
print("salida:", OUT)
